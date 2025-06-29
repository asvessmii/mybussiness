"""
Модуль для обработки различных типов документов
"""

import os
import re
import tempfile
from typing import Dict, Any, List
import logging

# Импорты для обработки документов
import PyPDF2
from docx import Document
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

import config

class DocumentProcessor:
    """Класс для обработки различных типов документов"""
    
    def __init__(self):
        # Настройка Tesseract OCR
        if hasattr(config, 'TESSERACT_CMD') and config.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_CMD
        
        self.supported_formats = config.ALLOWED_EXTENSIONS
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Основной метод обработки документа"""
        try:
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'error': 'Файл не найден'
                }
            
            # Определение типа файла
            file_extension = file_path.lower().split('.')[-1]
            
            if file_extension not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Неподдерживаемый формат файла: {file_extension}'
                }
            
            # Обработка в зависимости от типа файла
            if file_extension == 'pdf':
                result = self._process_pdf(file_path)
            elif file_extension == 'docx':
                result = self._process_docx(file_path)
            elif file_extension == 'txt':
                result = self._process_txt(file_path)
            else:
                return {
                    'success': False,
                    'error': f'Обработчик для {file_extension} не реализован'
                }
            
            if result['success']:
                # Очистка и нормализация текста
                cleaned_text = self._clean_text(result['text'])
                result['text'] = cleaned_text
                result['word_count'] = len(cleaned_text.split())
                result['char_count'] = len(cleaned_text)
            
            return result
            
        except Exception as e:
            logging.error(f"Ошибка обработки документа {file_path}: {e}")
            return {
                'success': False,
                'error': f'Ошибка обработки: {str(e)}'
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Обработка PDF файлов"""
        try:
            text_content = []
            pages_processed = 0
            
            # Попытка извлечения текста напрямую из PDF
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_content.append(page_text)
                            pages_processed += 1
                        else:
                            # Если текст не извлекается, используем OCR
                            ocr_text = self._extract_text_with_ocr(file_path, page_num)
                            if ocr_text:
                                text_content.append(ocr_text)
                                pages_processed += 1
                    except Exception as e:
                        logging.warning(f"Ошибка обработки страницы {page_num}: {e}")
                        continue
            
            if not text_content:
                return {
                    'success': False,
                    'error': 'Не удалось извлечь текст из PDF'
                }
            
            return {
                'success': True,
                'text': '\n\n'.join(text_content),
                'pages_processed': pages_processed,
                'format': 'pdf'
            }
            
        except Exception as e:
            logging.error(f"Ошибка обработки PDF {file_path}: {e}")
            return {
                'success': False,
                'error': f'Ошибка обработки PDF: {str(e)}'
            }
    
    def _extract_text_with_ocr(self, pdf_path: str, page_num: int) -> str:
        """Извлечение текста с помощью OCR"""
        try:
            # Конвертация страницы PDF в изображение
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(
                    pdf_path, 
                    first_page=page_num + 1,
                    last_page=page_num + 1,
                    output_folder=temp_dir,
                    dpi=200
                )
                
                if images:
                    # OCR для извлечения текста
                    text = pytesseract.image_to_string(images[0], lang='rus+eng')
                    return text
                
        except Exception as e:
            logging.warning(f"Ошибка OCR для страницы {page_num}: {e}")
            return ""
        
        return ""
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Обработка DOCX файлов"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # Извлечение текста из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                return {
                    'success': False,
                    'error': 'Документ не содержит текста'
                }
            
            return {
                'success': True,
                'text': '\n\n'.join(text_content),
                'pages_processed': 1,
                'format': 'docx'
            }
            
        except Exception as e:
            logging.error(f"Ошибка обработки DOCX {file_path}: {e}")
            return {
                'success': False,
                'error': f'Ошибка обработки DOCX: {str(e)}'
            }
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Обработка TXT файлов"""
        try:
            # Попытка определения кодировки
            encodings = ['utf-8', 'cp1251', 'iso-8859-1', 'utf-16']
            text_content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                return {
                    'success': False,
                    'error': 'Не удалось определить кодировку файла'
                }
            
            if not text_content.strip():
                return {
                    'success': False,
                    'error': 'Файл пустой'
                }
            
            return {
                'success': True,
                'text': text_content,
                'pages_processed': 1,
                'format': 'txt'
            }
            
        except Exception as e:
            logging.error(f"Ошибка обработки TXT {file_path}: {e}")
            return {
                'success': False,
                'error': f'Ошибка обработки TXT: {str(e)}'
            }
    
    def _clean_text(self, text: str) -> str:
        """Очистка и нормализация текста"""
        if not text:
            return ""
        
        # Удаление лишних пробелов и переносов строк
        text = re.sub(r'\s+', ' ', text)
        
        # Удаление специальных символов и управляющих символов
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f]', '', text)
        
        # Нормализация кавычек
        text = re.sub(r'[""„"«»]', '"', text)
        text = re.sub(r"[''‚']", "'", text)
        
        # Удаление повторяющихся знаков препинания
        text = re.sub(r'([.!?]){2,}', r'\1', text)
        text = re.sub(r'([,;:]){2,}', r'\1', text)
        
        # Удаление лишних пробелов вокруг знаков препинания
        text = re.sub(r'\s+([.!?,:;])', r'\1', text)
        text = re.sub(r'([.!?])\s*([A-ZА-Я])', r'\1 \2', text)
        
        # Удаление пустых строк
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if line.strip()]
        text = '\n'.join(cleaned_lines)
        
        return text.strip()
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Извлечение метаданных документа"""
        try:
            metadata = {
                'filename': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'file_extension': file_path.lower().split('.')[-1],
                'created_at': os.path.getctime(file_path),
                'modified_at': os.path.getmtime(file_path)
            }
            
            # Дополнительные метаданные для PDF
            if metadata['file_extension'] == 'pdf':
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        metadata['pages_count'] = len(pdf_reader.pages)
                        
                        if pdf_reader.metadata:
                            metadata['title'] = pdf_reader.metadata.get('/Title', '')
                            metadata['author'] = pdf_reader.metadata.get('/Author', '')
                            metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
                            metadata['creator'] = pdf_reader.metadata.get('/Creator', '')
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            logging.error(f"Ошибка извлечения метаданных {file_path}: {e}")
            return {}
    
    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """Валидация документа перед обработкой"""
        try:
            if not os.path.exists(file_path):
                return {'valid': False, 'error': 'Файл не существует'}
            
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return {'valid': False, 'error': 'Файл пустой'}
            
            if file_size > config.MAX_CONTENT_LENGTH:
                return {'valid': False, 'error': 'Файл слишком большой'}
            
            file_extension = file_path.lower().split('.')[-1]
            if file_extension not in self.supported_formats:
                return {'valid': False, 'error': f'Неподдерживаемый формат: {file_extension}'}
            
            return {'valid': True}
            
        except Exception as e:
            return {'valid': False, 'error': f'Ошибка валидации: {str(e)}'}